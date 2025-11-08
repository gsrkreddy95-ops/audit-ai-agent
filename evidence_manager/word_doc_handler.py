"""
Word Document Handler
Generates Word documents with audit evidence explanations and verifications
"""

import os
from pathlib import Path
from datetime import datetime
from typing import Optional, Dict
from rich.console import Console

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

console = Console()


class WordDocHandler:
    """
    Creates and manages Word documents for audit evidence
    """
    
    def __init__(self):
        if not DOCX_AVAILABLE:
            console.print("[yellow]⚠️  python-docx not installed. Word generation disabled.[/yellow]")
            console.print("[yellow]💡 Install with: pip install python-docx[/yellow]")
    
    def create_evidence_document(
        self,
        rfi_code: str,
        title: str,
        content: str,
        output_path: str,
        metadata: Optional[Dict] = None
    ) -> bool:
        """
        Create a Word document with evidence explanation
        
        Args:
            rfi_code: RFI code (e.g., "BCR-06.01")
            title: Document title
            content: Main content/explanation
            output_path: Where to save the document
            metadata: Optional metadata (author, date, etc.)
        
        Returns:
            True if successful
        """
        
        if not DOCX_AVAILABLE:
            console.print("[red]❌ python-docx not installed, cannot create Word document[/red]")
            return False
        
        try:
            console.print(f"[cyan]📝 Creating Word document for RFI {rfi_code}...[/cyan]")
            
            # Create document
            doc = Document()
            
            # Set document properties
            core_properties = doc.core_properties
            core_properties.author = metadata.get('author', 'Audit AI Agent') if metadata else 'Audit AI Agent'
            core_properties.title = title
            core_properties.subject = f"Audit Evidence for RFI {rfi_code}"
            core_properties.created = datetime.now()
            
            # Add header
            header = doc.sections[0].header
            header_para = header.paragraphs[0]
            header_para.text = f"RFI {rfi_code} - {title}"
            header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Add title
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(title)
            title_run.bold = True
            title_run.font.size = Pt(18)
            title_run.font.color.rgb = RGBColor(0, 51, 102)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph()  # Spacing
            
            # Add RFI code
            rfi_para = doc.add_paragraph()
            rfi_run = rfi_para.add_run(f"RFI Code: {rfi_code}")
            rfi_run.bold = True
            rfi_run.font.size = Pt(12)
            
            # Add timestamp
            timestamp_para = doc.add_paragraph()
            timestamp_run = timestamp_para.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S UTC')}")
            timestamp_run.font.size = Pt(10)
            timestamp_run.font.color.rgb = RGBColor(128, 128, 128)
            
            doc.add_paragraph()  # Spacing
            
            # Add metadata if provided
            if metadata:
                meta_table = doc.add_table(rows=len(metadata), cols=2)
                meta_table.style = 'Light Grid Accent 1'
                
                for i, (key, value) in enumerate(metadata.items()):
                    row = meta_table.rows[i]
                    row.cells[0].text = key.replace('_', ' ').title()
                    row.cells[1].text = str(value)
                
                doc.add_paragraph()  # Spacing
            
            # Add content heading
            content_heading = doc.add_heading('Evidence Explanation', level=1)
            
            # Add main content
            # Split by paragraphs
            for para_text in content.split('\n\n'):
                if para_text.strip():
                    # Check if it's a heading (starts with #)
                    if para_text.startswith('#'):
                        level = para_text.count('#', 0, 3)
                        text = para_text.lstrip('#').strip()
                        doc.add_heading(text, level=min(level, 3))
                    # Check if it's a bullet point
                    elif para_text.startswith('- ') or para_text.startswith('* '):
                        bullet_text = para_text.lstrip('- *').strip()
                        doc.add_paragraph(bullet_text, style='List Bullet')
                    # Check if it's a numbered list
                    elif para_text[0].isdigit() and para_text[1:3] in ['. ', ') ']:
                        list_text = para_text[3:].strip()
                        doc.add_paragraph(list_text, style='List Number')
                    else:
                        doc.add_paragraph(para_text)
            
            # Add footer
            footer = doc.sections[0].footer
            footer_para = footer.paragraphs[0]
            footer_para.text = f"Audit Evidence Collection | {datetime.now().strftime('%Y')}"
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Save document
            Path(output_path).parent.mkdir(parents=True, exist_ok=True)
            doc.save(output_path)
            
            console.print(f"[green]✅ Word document created: {output_path}[/green]")
            return True
            
        except Exception as e:
            console.print(f"[red]❌ Error creating Word document: {e}[/red]")
            import traceback
            traceback.print_exc()
            return False
    
    def create_verification_document(
        self,
        rfi_code: str,
        controls: list,
        output_path: str
    ) -> bool:
        """
        Create a Word document with control verification checklist
        
        Args:
            rfi_code: RFI code
            controls: List of controls with verification status
            output_path: Where to save
        
        Returns:
            True if successful
        """
        
        if not DOCX_AVAILABLE:
            console.print("[red]❌ python-docx not installed[/red]")
            return False
        
        try:
            console.print(f"[cyan]📝 Creating verification document for RFI {rfi_code}...[/cyan]")
            
            doc = Document()
            
            # Title
            title = f"Control Verification - RFI {rfi_code}"
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(title)
            title_run.bold = True
            title_run.font.size = Pt(16)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph()
            
            # Timestamp
            doc.add_paragraph(f"Verification Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S UTC')}")
            doc.add_paragraph()
            
            # Controls table
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Light Grid Accent 1'
            
            # Header row
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Control ID'
            header_cells[1].text = 'Description'
            header_cells[2].text = 'Status'
            header_cells[3].text = 'Notes'
            
            # Make header bold
            for cell in header_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            
            # Add controls
            for control in controls:
                row = table.add_row().cells
                row[0].text = control.get('id', 'N/A')
                row[1].text = control.get('description', 'N/A')
                
                status = control.get('status', 'Unknown')
                row[2].text = status
                
                # Color-code status
                if status.lower() in ['verified', 'pass', 'yes']:
                    for paragraph in row[2].paragraphs:
                        for run in paragraph.runs:
                            run.font.color.rgb = RGBColor(0, 128, 0)  # Green
                elif status.lower() in ['failed', 'fail', 'no']:
                    for paragraph in row[2].paragraphs:
                        for run in paragraph.runs:
                            run.font.color.rgb = RGBColor(255, 0, 0)  # Red
                else:
                    for paragraph in row[2].paragraphs:
                        for run in paragraph.runs:
                            run.font.color.rgb = RGBColor(255, 165, 0)  # Orange
                
                row[3].text = control.get('notes', '')
            
            # Save
            Path(output_path).parent.mkdir(parents=True, exist_ok=True)
            doc.save(output_path)
            
            console.print(f"[green]✅ Verification document created: {output_path}[/green]")
            return True
            
        except Exception as e:
            console.print(f"[red]❌ Error creating verification document: {e}[/red]")
            return False
    
    def update_existing_document(
        self,
        doc_path: str,
        new_content: str,
        append: bool = True
    ) -> bool:
        """
        Update an existing Word document
        
        Args:
            doc_path: Path to existing document
            new_content: Content to add/replace
            append: If True, append; if False, replace
        
        Returns:
            True if successful
        """
        
        if not DOCX_AVAILABLE:
            console.print("[red]❌ python-docx not installed[/red]")
            return False
        
        try:
            if not os.path.exists(doc_path):
                console.print(f"[red]❌ Document not found: {doc_path}[/red]")
                return False
            
            console.print(f"[cyan]📝 Updating document: {doc_path}...[/cyan]")
            
            doc = Document(doc_path)
            
            if not append:
                # Clear existing content
                for element in doc.element.body:
                    doc.element.body.remove(element)
            
            # Add update timestamp
            doc.add_paragraph()
            update_para = doc.add_paragraph()
            update_run = update_para.add_run(f"Updated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S UTC')}")
            update_run.font.size = Pt(10)
            update_run.font.color.rgb = RGBColor(128, 128, 128)
            
            # Add new content
            for para_text in new_content.split('\n\n'):
                if para_text.strip():
                    doc.add_paragraph(para_text)
            
            # Save
            doc.save(doc_path)
            
            console.print(f"[green]✅ Document updated: {doc_path}[/green]")
            return True
            
        except Exception as e:
            console.print(f"[red]❌ Error updating document: {e}[/red]")
            return False
    
    def extract_text_from_document(self, doc_path: str) -> Optional[str]:
        """
        Extract all text from a Word document
        
        Args:
            doc_path: Path to Word document
        
        Returns:
            Extracted text or None
        """
        
        if not DOCX_AVAILABLE:
            console.print("[red]❌ python-docx not installed[/red]")
            return None
        
        try:
            if not os.path.exists(doc_path):
                console.print(f"[red]❌ Document not found: {doc_path}[/red]")
                return None
            
            doc = Document(doc_path)
            
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text for cell in row.cells)
                    text_parts.append(row_text)
            
            full_text = '\n\n'.join(text_parts)
            
            console.print(f"[green]✅ Extracted {len(full_text)} characters from {doc_path}[/green]")
            return full_text
            
        except Exception as e:
            console.print(f"[red]❌ Error extracting text: {e}[/red]")
            return None
